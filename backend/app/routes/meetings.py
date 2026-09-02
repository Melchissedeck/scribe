import io
from datetime import date, datetime, time

from fastapi import APIRouter, Depends, HTTPException, Query, Response
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session, selectinload

from app.database import get_db
from app.dependencies import get_current_user
from app.models.action import Action
from app.models.recording import Recording
from app.models.speaker import Speaker
from app.models.transcript_segment import TranscriptSegment
from app.models.user import User
from app.schemas.action import ActionResponse
from app.schemas.recording import (
    AnonymizeResponse,
    DiarizedTranscriptResponse,
    MeetingDetailResponse,
    MeetingListItem,
    SegmentClassificationOut,
    SegmentClassificationResponse,
    SegmentOut,
    SpeakerOut,
    SpeakingTimeEntry,
    SpeakingTimeResponse,
    ThemeResponse,
    ThemeUpdate,
)
from app.services.anonymization_service import anonymize_recording
from app.services.llm_service import LLMService
from pdf_export_service import PDFExportService

router = APIRouter(prefix='/meetings', tags=['meetings'])

EXCERPT_LENGTH = 150

_STATUS_LABELS = {'todo': 'À faire', 'in_progress': 'En cours', 'done': 'Terminé'}


def _build_excerpt(summary: str | None) -> str | None:
    """Tronque un résumé à `EXCERPT_LENGTH` caractères pour l'affichage en liste.

    Args:
        summary: Résumé complet de la réunion, éventuellement absent.

    Returns:
        Le résumé tel quel s'il tient déjà dans la limite, une version
        tronquée suivie de '...' sinon, ou None si aucun résumé n'existe.
    """
    if not summary:
        return None
    if len(summary) <= EXCERPT_LENGTH:
        return summary
    return summary[:EXCERPT_LENGTH].rstrip() + '...'


@router.get('', response_model=list[MeetingListItem])
def list_meetings(
    theme: str | None = Query(default=None, description="Filtre sur le thème (recherche partielle, insensible à la casse)"),
    date_from: date | None = Query(default=None, description="Réunions à partir de cette date (incluse)"),
    date_to: date | None = Query(default=None, description="Réunions jusqu'à cette date (incluse)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Liste les réunions de l'utilisateur courant, avec filtres optionnels.

    Args:
        theme: Filtre sur le thème (recherche partielle, insensible à la casse).
        date_from: Ne renvoie que les réunions à partir de cette date (incluse).
        date_to: Ne renvoie que les réunions jusqu'à cette date (incluse).

    Returns:
        Liste des réunions correspondantes, triées par date décroissante.
    """
    query = db.query(Recording).filter(Recording.user_id == current_user.id)

    if theme:
        query = query.filter(Recording.theme.ilike(f'%{theme}%'))

    if date_from:
        query = query.filter(Recording.started_at >= datetime.combine(date_from, time.min))

    if date_to:
        query = query.filter(Recording.started_at <= datetime.combine(date_to, time.max))

    recordings = query.order_by(Recording.started_at.desc()).all()

    return [
        MeetingListItem(
            id=recording.id,
            theme=recording.theme,
            meeting_type='in_person' if recording.platform == 'dictaphone' else 'remote',
            date=recording.started_at,
            status=recording.status,
            summary_status=recording.summary_status,
            summary_excerpt=_build_excerpt(recording.summary),
            duration_minutes=(
                (recording.stopped_at - recording.started_at).total_seconds() / 60
                if recording.stopped_at
                else 0
            ),
        )
        for recording in recordings
    ]


@router.patch('/{meeting_id}/theme', response_model=ThemeResponse)
def update_meeting_theme(
    meeting_id: int,
    payload: ThemeUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Met à jour le thème d'une réunion appartenant à l'utilisateur courant.

    Args:
        meeting_id: Identifiant de la réunion.
        payload: Nouveau thème à appliquer (vide ou blanc efface le thème).

    Returns:
        L'identifiant de la réunion et son thème mis à jour.

    Raises:
        HTTPException: 404 si la réunion est introuvable ou n'appartient pas à l'utilisateur.
    """
    recording = db.query(Recording).filter(
        Recording.id == meeting_id, Recording.user_id == current_user.id
    ).first()
    if not recording:
        raise HTTPException(status_code=404, detail='Réunion introuvable.')

    recording.theme = payload.theme.strip() if payload.theme and payload.theme.strip() else None
    db.commit()
    db.refresh(recording)

    return ThemeResponse(recording_id=recording.id, theme=recording.theme)


@router.get('/{meeting_id}/speaking-time', response_model=SpeakingTimeResponse)
def get_speaking_time(
    meeting_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Calcule le temps de parole de chaque locuteur d'une réunion.

    Args:
        meeting_id: Identifiant de la réunion.

    Returns:
        Le temps de parole (en secondes et en pourcentage) par locuteur,
        du plus bavard au moins bavard.

    Raises:
        HTTPException: 404 si la réunion est introuvable ou n'appartient pas à l'utilisateur.
    """
    recording = db.query(Recording).filter(
        Recording.id == meeting_id, Recording.user_id == current_user.id
    ).first()
    if not recording:
        raise HTTPException(status_code=404, detail='Réunion introuvable.')

    segments = db.query(TranscriptSegment).filter(
        TranscriptSegment.recording_id == meeting_id
    ).all()

    durations: dict[str, float] = {}
    for seg in segments:
        duration = max(0.0, seg.end - seg.start)
        durations[seg.speaker] = durations.get(seg.speaker, 0.0) + duration

    total = sum(durations.values())

    if total == 0:
        return SpeakingTimeResponse(meeting_id=meeting_id, entries=[])

    entries = [
        SpeakingTimeEntry(
            speaker=speaker,
            seconds=round(seconds, 1),
            percentage=round(seconds / total * 100, 1),
        )
        for speaker, seconds in sorted(durations.items(), key=lambda x: -x[1])
    ]

    return SpeakingTimeResponse(meeting_id=meeting_id, entries=entries)


@router.get('/{meeting_id}/diarized-transcript', response_model=DiarizedTranscriptResponse)
def get_diarized_transcript(
    meeting_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Récupère la transcription diarisée d'une réunion, segment par segment.

    Args:
        meeting_id: Identifiant de la réunion.

    Returns:
        La liste des segments diarisés (locuteur et texte).

    Raises:
        HTTPException: 404 si la réunion est introuvable, n'appartient pas
            à l'utilisateur, ou si aucune diarisation n'est disponible.
    """
    recording = (
        db.query(Recording)
        .filter(Recording.id == meeting_id, Recording.user_id == current_user.id)
        .first()
    )
    if not recording:
        raise HTTPException(status_code=404, detail='Réunion introuvable.')

    segments = (
        db.query(TranscriptSegment)
        .filter(TranscriptSegment.recording_id == meeting_id)
        .order_by(TranscriptSegment.start)
        .all()
    )

    if not segments:
        raise HTTPException(
            status_code=404,
            detail='Aucune diarisation disponible pour cette réunion.',
        )

    return DiarizedTranscriptResponse(
        meeting_id=meeting_id,
        segments=[
            SegmentOut(speaker_name=seg.speaker, text=seg.text)
            for seg in segments
        ],
    )


@router.post('/{meeting_id}/classify-segments', response_model=SegmentClassificationResponse)
def classify_segments(
    meeting_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Classifie chaque segment diarisé d'une réunion via le LLM (ton, thème, urgence).

    Les classifications obtenues sont enregistrées sur les segments en base.

    Args:
        meeting_id: Identifiant de la réunion.

    Returns:
        Les segments de la réunion, avec leur classification.

    Raises:
        HTTPException: 404 si la réunion est introuvable ou n'appartient
            pas à l'utilisateur ; 400 si aucune transcription diarisée
            n'est disponible ; 502 si le service de classification est
            momentanément indisponible.
    """
    recording = (
        db.query(Recording)
        .filter(Recording.id == meeting_id, Recording.user_id == current_user.id)
        .first()
    )
    if not recording:
        raise HTTPException(status_code=404, detail='Réunion introuvable.')

    segments = (
        db.query(TranscriptSegment)
        .filter(TranscriptSegment.recording_id == meeting_id)
        .order_by(TranscriptSegment.start)
        .all()
    )
    if not segments:
        raise HTTPException(
            status_code=400,
            detail='Aucune transcription diarisée disponible pour cette réunion.',
        )

    llm_service = LLMService()
    result = llm_service.classify_segments([str(seg.text) for seg in segments])
    if result is None:
        raise HTTPException(
            status_code=502,
            detail='Le service de classification est momentanément indisponible.',
        )

    classifications_by_index = {c.index: c for c in result.classifications}
    for i, seg in enumerate(segments):
        classification = classifications_by_index.get(i)
        if classification is None:
            continue
        seg.tone = classification.tone
        seg.theme = classification.theme
        seg.urgency = classification.urgency

    db.commit()

    return SegmentClassificationResponse(
        meeting_id=recording.id,
        segments=[
            SegmentClassificationOut(
                id=seg.id,
                speaker_name=seg.speaker,
                text=seg.text,
                start=seg.start,
                tone=seg.tone,
                theme=seg.theme,
                urgency=seg.urgency,
            )
            for seg in segments
        ],
    )


@router.get('/{meeting_id}/segments-classification', response_model=SegmentClassificationResponse)
def get_segments_classification(
    meeting_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Récupère la classification déjà calculée des segments diarisés d'une réunion.

    Args:
        meeting_id: Identifiant de la réunion.

    Returns:
        Les segments de la réunion, avec leur classification.

    Raises:
        HTTPException: 404 si la réunion est introuvable, n'appartient pas
            à l'utilisateur, ou si aucune diarisation n'est disponible.
    """
    recording = (
        db.query(Recording)
        .filter(Recording.id == meeting_id, Recording.user_id == current_user.id)
        .first()
    )
    if not recording:
        raise HTTPException(status_code=404, detail='Réunion introuvable.')

    segments = (
        db.query(TranscriptSegment)
        .filter(TranscriptSegment.recording_id == meeting_id)
        .order_by(TranscriptSegment.start)
        .all()
    )
    if not segments:
        raise HTTPException(
            status_code=404,
            detail='Aucune diarisation disponible pour cette réunion.',
        )

    return SegmentClassificationResponse(
        meeting_id=recording.id,
        segments=[
            SegmentClassificationOut(
                id=seg.id,
                speaker_name=seg.speaker,
                text=seg.text,
                start=seg.start,
                tone=seg.tone,
                theme=seg.theme,
                urgency=seg.urgency,
            )
            for seg in segments
        ],
    )


@router.post('/{meeting_id}/anonymize', response_model=AnonymizeResponse)
def anonymize_meeting(
    meeting_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Remplace les noms des locuteurs d'une réunion par un libellé générique.

    Action irréversible (US-69, droit à l'effacement RGPD) : le nom
    d'origine n'est conservé nulle part une fois l'opération effectuée.
    """
    recording = (
        db.query(Recording)
        .filter(Recording.id == meeting_id, Recording.user_id == current_user.id)
        .first()
    )
    if not recording:
        raise HTTPException(status_code=404, detail='Réunion introuvable.')

    anonymize_recording(db, recording.id)

    segments = (
        db.query(TranscriptSegment)
        .filter(TranscriptSegment.recording_id == meeting_id)
        .order_by(TranscriptSegment.start)
        .all()
    )

    return AnonymizeResponse(
        recording_id=recording.id,
        segments=[
            SegmentOut(speaker_name=seg.speaker, text=seg.text, start=seg.start)
            for seg in segments
        ],
    )


@router.get('/{meeting_id}/details', response_model=MeetingDetailResponse)
def get_meeting_details(
    meeting_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Récupère le détail complet d'une réunion (locuteurs, segments, actions).

    Args:
        meeting_id: Identifiant de la réunion.

    Returns:
        Les informations détaillées de la réunion.

    Raises:
        HTTPException: 404 si la réunion est introuvable ou n'appartient pas à l'utilisateur.
    """
    recording = (
        db.query(Recording)
        .filter(Recording.id == meeting_id, Recording.user_id == current_user.id)
        .options(
            selectinload(Recording.speakers),
            selectinload(Recording.segments),
            selectinload(Recording.actions),
        )
        .first()
    )
    if not recording:
        raise HTTPException(status_code=404, detail='Réunion introuvable.')

    segments = sorted(recording.segments, key=lambda seg: seg.start)

    return MeetingDetailResponse(
        id=recording.id,
        theme=recording.theme,
        meeting_type='in_person' if recording.platform == 'dictaphone' else 'remote',
        status=recording.status,
        platform=recording.platform,
        diarization_status=recording.diarization_status,
        started_at=recording.started_at,
        stopped_at=recording.stopped_at,
        summary=recording.summary,
        speakers=[SpeakerOut.model_validate(speaker) for speaker in recording.speakers],
        segments=[SegmentOut(speaker_name=seg.speaker, text=seg.text, start=seg.start) for seg in segments],
        actions=[ActionResponse.model_validate(action) for action in recording.actions],
    )


@router.get('/{meeting_id}/export-docx')
def export_docx(
    meeting_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> StreamingResponse:
    """Génère et retourne le compte-rendu d'une réunion au format Word (.docx).

    Le document inclut le titre, la date, le résumé (converti depuis le
    Markdown) et le plan d'action sous forme de tableau.

    Args:
        meeting_id: Identifiant de la réunion.

    Returns:
        Le fichier .docx généré, en réponse de type flux (streaming).

    Raises:
        HTTPException: 404 si la réunion est introuvable ou n'appartient
            pas à l'utilisateur ; 503 si le module de génération Word
            n'est pas disponible sur le serveur ; 500 si la génération du
            fichier échoue.
    """
    recording = (
        db.query(Recording)
        .filter(Recording.id == meeting_id, Recording.user_id == current_user.id)
        .options(selectinload(Recording.actions).selectinload(Action.speaker))
        .first()
    )
    if not recording:
        raise HTTPException(status_code=404, detail='Réunion introuvable.')

    try:
        import re as _re

        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise HTTPException(status_code=503, detail='Module Word non disponible sur ce serveur.') from exc

    def _add_md_paragraph(doc, text: str, style: str = 'Normal'):
        """Ajoute un paragraphe au document Word en convertissant le **gras** Markdown.

        Args:
            doc: Document Word en cours de construction.
            text: Texte du paragraphe, pouvant contenir des séquences `**gras**`.
            style: Style de paragraphe Word à appliquer.

        Returns:
            Le paragraphe ajouté au document.
        """
        parts = _re.split(r'\*\*(.+?)\*\*', text)
        p = doc.add_paragraph(style=style)
        for i, part in enumerate(parts):
            run = p.add_run(part)
            if i % 2 == 1:
                run.bold = True
        return p

    def _add_md_summary(doc, summary: str):
        """Convertit un résumé au format Markdown en titres et paragraphes Word.

        Interprète les titres (`#`, `##`, `###`) et les listes à puces
        (`-`, `*`) ligne par ligne ; toute autre ligne devient un
        paragraphe normal (voir `_add_md_paragraph`).

        Args:
            doc: Document Word en cours de construction.
            summary: Résumé au format Markdown à convertir.
        """
        for line in summary.split('\n'):
            line = line.rstrip()
            if not line:
                doc.add_paragraph('')
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], 2)
            elif line.startswith('- ') or line.startswith('* '):
                _add_md_paragraph(doc, line[2:], 'List Bullet')
            else:
                _add_md_paragraph(doc, line)

    try:
        doc = Document()

        title = str(recording.theme).strip() if recording.theme else 'Réunion sans titre'
        doc.add_heading(title, 0)

        if recording.started_at:
            date_str = recording.started_at.strftime('%d/%m/%Y à %H:%M')
            meta = doc.add_paragraph(f'Date : {date_str}')
            meta.runs[0].font.size = Pt(11)

        doc.add_heading('Résumé', 1)
        summary_text = str(recording.summary).strip() if recording.summary else ''
        if summary_text:
            _add_md_summary(doc, summary_text)
        else:
            doc.add_paragraph('Aucun résumé disponible.')

        actions = list(recording.actions)
        if actions:
            doc.add_heading("Plan d'action", 1)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Description'
            hdr[1].text = 'Responsable'
            hdr[2].text = 'Statut'
            hdr[3].text = 'Échéance'
            for action in actions:
                row = table.add_row().cells
                speaker = action.speaker
                responsable = speaker.provisional_name if speaker else '—'
                row[0].text = str(action.description)
                row[1].text = responsable
                row[2].text = _STATUS_LABELS.get(action.status, action.status)
                row[3].text = action.due_date.strftime('%d/%m/%Y') if action.due_date else '—'

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        safe = ''.join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
        filename = f"{safe}.docx" if safe else f"compte-rendu-{meeting_id}.docx"

        return StreamingResponse(
            buffer,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'},
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail='Erreur lors de la génération du fichier Word.') from exc


@router.delete('/{meeting_id}', status_code=204)
def delete_meeting(
    meeting_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> Response:
    """Supprime définitivement une réunion et toutes ses données associées.

    Supprime en cascade les actions, segments de transcription et
    locuteurs liés à la réunion avant de supprimer la réunion elle-même.

    Args:
        meeting_id: Identifiant de la réunion à supprimer.

    Raises:
        HTTPException: 404 si la réunion est introuvable ou n'appartient pas à l'utilisateur.
    """
    recording = db.query(Recording).filter(
        Recording.id == meeting_id, Recording.user_id == current_user.id
    ).first()
    if not recording:
        raise HTTPException(status_code=404, detail='Réunion introuvable.')

    db.query(Action).filter(Action.recording_id == meeting_id).delete()
    db.query(TranscriptSegment).filter(TranscriptSegment.recording_id == meeting_id).delete()
    db.query(Speaker).filter(Speaker.recording_id == meeting_id).delete()
    db.delete(recording)
    db.commit()
    return Response(status_code=204)


@router.get('/{meeting_id}/export-pdf')
def export_meeting_pdf(
    meeting_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Génère et retourne le compte-rendu d'une réunion au format PDF.

    Args:
        meeting_id: Identifiant de la réunion.

    Returns:
        Le fichier PDF généré, en pièce jointe.

    Raises:
        HTTPException: 404 si la réunion est introuvable ou n'appartient pas à l'utilisateur.
    """
    recording = (
        db.query(Recording)
        .filter(Recording.id == meeting_id, Recording.user_id == current_user.id)
        .options(
            selectinload(Recording.actions).selectinload(Action.speaker),
        )
        .first()
    )
    if not recording:
        raise HTTPException(status_code=404, detail='Réunion introuvable.')

    pdf_bytes = PDFExportService().generate_pdf(recording)

    return Response(
        content=pdf_bytes,
        media_type='application/pdf',
        headers={'Content-Disposition': f'attachment; filename="compte-rendu-{meeting_id}.pdf"'},
    )
